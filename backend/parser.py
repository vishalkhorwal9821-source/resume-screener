import io
import re

import docx
import fitz

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s@.,+\-/#]', ' ', text)
    return text.strip()

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        text_parts = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text("text") or "")
        text = " ".join(text_parts)
        return clean_text(text)
    except Exception:
        return ""

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return clean_text(" ".join(paragraphs))
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""
